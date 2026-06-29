import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def create_report():
    doc = docx.Document()

    # Set Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Styles styling
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Arial'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Helper function for headings
    def add_heading_1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x0D, 0x94, 0x88) # Teal-600
        return p

    def add_heading_2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x0F, 0x76, 0x6E) # Teal-700
        return p

    def set_cell_background(cell, fill_hex):
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
        cell._tc.get_or_add_tcPr().append(shading_elm)

    # Title Banner
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(20)
    title_p.paragraph_format.space_after = Pt(6)
    title_run = title_p.add_run("SMART PRODUCT SEARCH & FILTER")
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x0D, 0x94, 0x88)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(24)
    sub_run = sub_p.add_run("Feature & Implementation Report:\nAI Multimodal Visual Search Choice Modal & UI Modernization")
    sub_run.font.size = Pt(14)
    sub_run.font.bold = True
    sub_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Section 1: Overview
    add_heading_1("1. Executive Summary & Feature Overview")
    p = doc.add_paragraph(
        "As part of modernizing the Smart Product Search and Filter web application for Sharadha Stores, "
        "a major architectural enhancement was made to the AI-powered product discovery engine. "
        "To provide a frictionless and versatile user experience, an interactive glassmorphic Visual Search Choice Modal "
        "was implemented directly into the primary navigation bar. This feature empowers customers to search for traditional Indian "
        "groceries, sweets, pickles, and snacks using visual AI analysis through two distinct input methods:"
    )
    p.paragraph_format.space_after = Pt(8)

    bullet1 = doc.add_paragraph(style='List Bullet')
    r1 = bullet1.add_run("Option A: Upload / Scan Photo (📁): ")
    r1.bold = True
    bullet1.add_run("Allows users to upload an image file from their local device storage or capture a direct photograph using their mobile phone or laptop camera.")

    bullet2 = doc.add_paragraph(style='List Bullet')
    r2 = bullet2.add_run("Option B: Paste Image Web Link (🔗): ")
    r2.bold = True
    bullet2.add_run("Allows users to paste any public web URL linking to an image of a dish or food item. The system fetches the remote image buffer dynamically and processes it seamlessly.")

    # Section 2: Architecture & Workflow
    add_heading_1("2. Technical Architecture & AI Data Flow")
    p = doc.add_paragraph(
        "The Multimodal Visual Search workflow bridges the React (Vite) frontend with the Express.js serverless backend "
        "and Google Gemini Vision AI (gemini-2.5-flash). The step-by-step technical execution is detailed below:"
    )
    p.paragraph_format.space_after = Pt(12)

    # Table for Workflow Steps
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Step"
    hdr_cells[1].text = "Layer / Component"
    hdr_cells[2].text = "Technical Implementation Details"

    set_cell_background(hdr_cells[0], "0D9488")
    set_cell_background(hdr_cells[1], "0D9488")
    set_cell_background(hdr_cells[2], "0D9488")
    for cell in hdr_cells:
        for p_cell in cell.paragraphs:
            for run in p_cell.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    workflow_data = [
        ("1. User Trigger", "Frontend (Navbar.jsx)", "User clicks the Scan/Camera button inside the search bar. The state showScanModal is set to true, triggering a sleek, responsive modal overlay."),
        ("2. Input Selection", "Choice Modal UI", "User selects either 'Upload Photo' (triggering a hidden HTML5 file input) or 'Paste Link' (displaying a URL text input form with loading indicator)."),
        ("3. Payload Transmission", "REST API Client", "For file uploads, data is sent via multipart/form-data. For web URLs, a JSON payload { image_url: '...' } is transmitted to POST /api/visual-search."),
        ("4. Backend Buffer Processing", "Express.js (aiRoutes.js)", "If a file is received, it is read into memory and converted to Base64. If an image_url is received, the server fetches the remote blob via native fetch() and converts the ArrayBuffer to Base64 inline data."),
        ("5. Vision AI Query", "Google Gemini 2.5 Flash", "The Base64 inline data is passed along with a prompt instructing the AI to act as an Indian food identification expert and return only the exact dish name (1-3 words)."),
        ("6. Instant Filtering", "React State & Navigation", "The identified search term (e.g., 'Mango Pickle' or 'Samosa') is returned to the client, triggering an automatic product grid filter and navigating the user to the Shop page.")
    ]

    for step, layer, desc in workflow_data:
        row_cells = table.add_row().cells
        row_cells[0].text = step
        row_cells[1].text = layer
        row_cells[2].text = desc
        for cell in row_cells:
            for p_cell in cell.paragraphs:
                p_cell.paragraph_format.space_after = Pt(4)

    # Section 3: UI & UX Modernization
    add_heading_1("3. UI/UX Refinement & Authenticity Cleanup")
    p = doc.add_paragraph(
        "To elevate the overall visual aesthetic and ensure a professional, production-grade presentation, "
        "several mock and demo artifacts were purged across the codebase:"
    )
    
    cleanups = [
        ("Removal of Demo Credentials Banners: ", "Eliminated the bright green 'Demo account available' popup box on the Customer Login page and the 'Demo password: 1234' hint on the Admin Login page."),
        ("Elimination of False Testimonials & Ratings: ", "Removed hardcoded 4.8/5 average customer rating statistics, unverified testimonial carousels, and inflated trust metrics to maintain 100% data authenticity."),
        ("Footer Cleanups: ", "Replaced generic copyright claims with a clean, localized branding message ('Made with ❤️ in India for Sharadha Stores').")
    ]
    for title, desc in cleanups:
        bp = doc.add_paragraph(style='List Bullet')
        r = bp.add_run(title)
        r.bold = True
        bp.add_run(desc)

    # Section 4: Screenshot Placeholders
    add_heading_1("4. Verification & Screenshot Annexure")
    p = doc.add_paragraph("Insert screenshots below to document visual proof of the implementation in your final report:")
    
    boxes = [
        ("Figure 1: Visual Search Choice Modal (Upload vs. Paste Link Option)", "Insert screenshot showing the modal popup over the homepage when the scan icon is clicked."),
        ("Figure 2: URL Pasting Mode in Action", "Insert screenshot showing an image web link pasted into the modal input field before clicking search."),
        ("Figure 3: Clean Customer Login Page", "Insert screenshot demonstrating the modernized login page completely devoid of demo banners.")
    ]
    
    for cap, note in boxes:
        add_heading_2(cap)
        box_p = doc.add_paragraph()
        box_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        box_p.paragraph_format.space_before = Pt(10)
        box_p.paragraph_format.space_after = Pt(10)
        run = box_p.add_run(f"[ {note} ]")
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.save("Visual_Search_Choice_Modal_Documentation.docx")
    print("Successfully created Visual_Search_Choice_Modal_Documentation.docx")

if __name__ == "__main__":
    create_report()
